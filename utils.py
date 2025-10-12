# utils.py
import os
import requests
import docx
import PyPDF2
from flask import current_app, url_for, flash, render_template_string
from flask_mail import Message
from itsdangerous import URLSafeTimedSerializer
from werkzeug.security import generate_password_hash, check_password_hash

# optional OCR support
try:
    from PIL import Image
    import pytesseract
    from io import BytesIO
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

MAX_AI_RESPONSE_CHARS = 500  # max characters for concise AI answers
MAX_OCR_CHARS = 1000         # max chars extracted from images


def send_verification_email(user):
    from extensions import mail

    serializer = URLSafeTimedSerializer(current_app.config['SECRET_KEY'])
    token = serializer.dumps(user.email, salt='email-confirm')
    confirm_url = url_for('app_routes.confirm_email', token=token, _external=True)

    subject = "Confirm Your Mirai Account"
    sender = current_app.config['MAIL_USERNAME']
    recipient = user.email

    text_body = f"""Hi {user.email},

Please confirm your Mirai account by clicking the link below:

{confirm_url}

This link will expire in 1 hour.
"""
    html_body = render_template_string("""
    <h2>Welcome to Mirai!</h2>
    <p>Hi {{ email }},</p>
    <p>Click the button below to confirm your account:</p>
    <a href="{{ confirm_url }}" style="display:inline-block;background:#4CAF50;color:#fff;padding:10px 20px;text-decoration:none;border-radius:5px;">Confirm Email</a>
    <p>If the button doesn't work, use this link:</p>
    <p>{{ confirm_url }}</p>
    """, email=user.email, confirm_url=confirm_url)

    msg = Message(subject=subject, sender=sender, recipients=[recipient])
    msg.body = text_body
    msg.html = html_body

    try:
        mail.send(msg)
    except Exception as e:
        print(f"❌ Failed to send email: {e}")
        flash("⚠️ Failed to send confirmation email.", "danger")


def hash_password(password):
    return generate_password_hash(password)


def verify_password(password, hashed):
    return check_password_hash(hashed, password)


def read_stored_file_content(attachment, max_chars=MAX_OCR_CHARS):
    """
    Read a stored file (PDF, DOCX, TXT, or image) from server.
    For images:
        - Uses OCR if available.
        - If OCR fails, returns a short description.
    Returns a string to include in AI prompt.
    """
    upload_dir = os.path.join(current_app.instance_path, "uploads")
    stored_name = getattr(attachment, "path", None) or getattr(attachment, "stored_name", None)
    if not stored_name:
        return ""

    file_path = os.path.join(upload_dir, stored_name)
    filename = getattr(attachment, "filename", stored_name)
    ctype = (getattr(attachment, "content_type", "") or "").lower()

    if not os.path.exists(file_path):
        return f"[Attachment: {filename}] (file missing)"

    try:
        # TXT
        if ctype.startswith("text") or filename.lower().endswith(".txt"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as fh:
                return fh.read(max_chars)

        # PDF
        if filename.lower().endswith(".pdf") or ctype == "application/pdf":
            try:
                with open(file_path, "rb") as fh:
                    reader = PyPDF2.PdfReader(fh)
                    text = " ".join((page.extract_text() or "") for page in reader.pages)
                return text[:max_chars]
            except Exception:
                return f"[PDF: {filename}] (unable to parse)"

        # DOCX
        if filename.lower().endswith(".docx") or ctype in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",):
            try:
                doc = docx.Document(file_path)
                text = "\n".join(para.text for para in doc.paragraphs)
                return text[:max_chars]
            except Exception:
                return f"[DOCX: {filename}] (unable to parse)"

        # Images
        image_exts = (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp", ".gif")
        if filename.lower().endswith(image_exts) or ctype.startswith("image"):
            if OCR_AVAILABLE:
                try:
                    img = Image.open(file_path)
                    if img.mode not in ("RGB", "L"):
                        img = img.convert("RGB")
                    ocr_text = pytesseract.image_to_string(img).strip()
                    if ocr_text:
                        return ocr_text[:max_chars]
                except Exception:
                    pass
            # fallback description
            return f"[Image: {filename}] Description: An image is attached. Possibly contains objects or scenes. AI should consider this in the response."

        return f"[Attachment: {filename}] (unsupported type)"

    except Exception:
        return f"[Attachment: {filename}] (unreadable)"


def generate_response(user_msg, username, memory=None, attachment=None):
    """
    Generates a short response from Mirai AI.
    - attachment: optional Attachment object; content will be included.
    """
    if memory is None:
        memory = []

    messages = memory.copy()
    system_prompt = f"""You are Mirai, assistant for {username}.
Provide short, concise, and clear answers (max {MAX_AI_RESPONSE_CHARS} characters).
Be polite and safe. Incorporate any relevant info from attachments."""
    messages.insert(0, {"role": "system", "content": system_prompt})

    user_content = user_msg
    if attachment:
        ocr_text = read_stored_file_content(attachment)
        if ocr_text:
            user_content += f"\n\n[Attachment content]: {ocr_text}"

    messages.append({"role": "user", "content": user_content})

    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY')}",
                "Content-Type": "application/json"
            },
            json={
                "model": "meta-llama/llama-3-8b-instruct",
                "messages": messages,
                "max_tokens": MAX_AI_RESPONSE_CHARS // 4
            },
            timeout=30
        )
        data = response.json()
        if "choices" in data and len(data["choices"]) > 0:
            return data["choices"][0]["message"]["content"][:MAX_AI_RESPONSE_CHARS]
        return "⚠️ No response from AI."
    except Exception as e:
        current_app.logger.exception("AI error: %s", e)
        return "⚠️ Error contacting AI."


def generate_chat_title(prompt, username="User"):
    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY')}",
                "Content-Type": "application/json"
            },
            json={
                "model": "meta-llama/llama-4-maverick",
                "messages": [
                    {"role": "system", "content": "Generate a short title (max 3 words) for this chat. Give only title."},
                    {"role": "user", "content": prompt}
                ]
            },
            timeout=15
        )
        data = response.json()
        if "choices" in data and len(data["choices"]) > 0:
            return data["choices"][0]["message"]["content"].strip()
        return "Untitled Chat"
    except Exception:
        return "Untitled Chat"


def read_file_content(file_storage, max_chars=MAX_OCR_CHARS):
    """Read uploaded FileStorage object (TXT, PDF, DOCX, Image)"""
    filename = (getattr(file_storage, "filename", "") or "").lower()
    ctype = (getattr(file_storage, "mimetype", "") or "").lower()
    try:
        # TXT
        if filename.endswith(".txt") or ctype.startswith("text"):
            file_storage.stream.seek(0)
            text = file_storage.read().decode("utf-8", errors="ignore")
            file_storage.stream.seek(0)
            return text[:max_chars]

        # PDF
        if filename.endswith(".pdf") or ctype == "application/pdf":
            file_storage.stream.seek(0)
            reader = PyPDF2.PdfReader(file_storage)
            text = " ".join((page.extract_text() or "") for page in reader.pages)
            file_storage.stream.seek(0)
            return text[:max_chars]

        # DOCX
        if filename.endswith(".docx") or ctype in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",):
            file_storage.stream.seek(0)
            doc = docx.Document(file_storage)
            text = "\n".join(para.text for para in doc.paragraphs)
            file_storage.stream.seek(0)
            return text[:max_chars]

        # Image
        image_exts = (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp", ".gif")
        if filename.endswith(image_exts) or ctype.startswith("image"):
            if OCR_AVAILABLE:
                file_storage.stream.seek(0)
                from io import BytesIO
                img = Image.open(BytesIO(file_storage.read()))
                file_storage.stream.seek(0)
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")
                text = pytesseract.image_to_string(img).strip()
                return text[:max_chars]
            return ""
    except Exception:
        return "[Unreadable content]"
